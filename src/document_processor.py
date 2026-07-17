from docx import Document


class DocumentProcessor:

    def __init__(self, detector, mapper):

        self.detector = detector
        self.mapper = mapper

    def process_document(self, input_path: str, output_path: str):

        document = Document(input_path)

        self._process_paragraphs(document)

        self._process_tables(document)

        document.save(output_path)

        print("\nDocument redacted successfully.")
        print(f"Saved to: {output_path}")

    def _process_paragraphs(self, document):

        for paragraph in document.paragraphs:

            self._process_paragraph(paragraph)

    def _process_tables(self, document):

        for table in document.tables:

            for row in table.rows:

                for cell in row.cells:

                    for paragraph in cell.paragraphs:

                        self._process_paragraph(paragraph)

    def _process_paragraph(self, paragraph):

        if not paragraph.text.strip():

            return

        entities = self.detector.detect(paragraph.text)

        if not entities:

            return

        redacted_text = self._replace_entities(
            paragraph.text,
            entities
        )

        # Replace text while preserving the paragraph.
        # Formatting inside runs may be simplified, but
        # the overall document structure remains intact.

        if paragraph.runs:

            paragraph.runs[0].text = redacted_text

            for run in paragraph.runs[1:]:

                run.text = ""

        else:

            paragraph.text = redacted_text

    def _replace_entities(self, text, entities):

        entities.sort(
            key=lambda entity: entity.start,
            reverse=True
        )

        for entity in entities:

            fake_value = self.mapper.get_fake_value(
                entity.text,
                entity.entity_type
            )

            text = (
                text[:entity.start]
                + fake_value
                + text[entity.end:]
            )

        return text