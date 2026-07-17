import json
from pathlib import Path

from docx import Document

from src.detector import PIIDetector


class Evaluator:

    def __init__(self):
        self.detector = PIIDetector()

    def load_document(self, path):

        doc = Document(path)

        text = ""

        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"

        return text

    def load_ground_truth(self, path):

        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)["entities"]

    def evaluate(self, document_path, ground_truth_path):

        document = self.load_document(document_path)

        predictions = self.detector.detect(document)

        ground_truth = self.load_ground_truth(ground_truth_path)

        matched_predictions = set()
        matched_truth = set()

        tp = 0

        # Match by exact text + entity type
        for gt_index, gt in enumerate(ground_truth):

            gt_text = gt["text"].strip().lower()
            gt_type = gt["type"].upper()

            for pred_index, pred in enumerate(predictions):

                if pred_index in matched_predictions:
                    continue

                pred_text = pred.text.strip().lower()
                pred_type = pred.entity_type.upper()

                if pred_text == gt_text and pred_type == gt_type:

                    tp += 1

                    matched_predictions.add(pred_index)
                    matched_truth.add(gt_index)

                    break

        fp = len(predictions) - len(matched_predictions)

        fn = len(ground_truth) - len(matched_truth)

        precision = tp / (tp + fp) if (tp + fp) else 0

        recall = tp / (tp + fn) if (tp + fn) else 0

        f1 = (
            2 * precision * recall / (precision + recall)
            if (precision + recall)
            else 0
        )

        # Approximate accuracy
        accuracy = tp / (tp + fp + fn) if (tp + fp + fn) else 0

        report = f"""
==========================================================
PII REDACTION EVALUATION REPORT
==========================================================

Ground Truth Entities : {len(ground_truth)}
Detected Entities     : {len(predictions)}

True Positives  : {tp}
False Positives : {fp}
False Negatives : {fn}

----------------------------------------------------------

Accuracy  : {accuracy:.3f}

Precision : {precision:.3f}

Recall    : {recall:.3f}

F1 Score  : {f1:.3f}

==========================================================
"""

        print(report)

        Path("evaluation").mkdir(exist_ok=True)

        with open(
            "evaluation/evaluation_report.txt",
            "w",
            encoding="utf-8",
        ) as f:

            f.write(report)

        print(
            "Evaluation report saved to evaluation/evaluation_report.txt"
        )


if __name__ == "__main__":

    evaluator = Evaluator()

    evaluator.evaluate(
        "input/Red Herring Prospectus.docx",
        "evaluation/ground_truth.json",
    )